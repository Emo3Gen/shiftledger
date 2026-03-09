#!/usr/bin/env python3
"""
ЭТАП 2: Извлечение базы знаний и паттернов ответов
===================================================

1. Парсинг документов (Правила, Договор оферты)
2. Извлечение паттернов ответов операторов из реальных диалогов
3. Кластеризация типичных вопросов клиентов
4. Формирование knowledge_base.json
"""

import json
import os
import re
from collections import Counter, defaultdict
from datetime import datetime

try:
    import docx
except ImportError:
    import subprocess
    subprocess.run(["pip", "install", "python-docx", "--break-system-packages", "-q"])
    import docx

INPUT_DIR = "/mnt/user-data/uploads"
OUTPUT_DIR = "/home/claude/shiftledger/output"
KB_DIR = "/home/claude/shiftledger/knowledge_base"

os.makedirs(KB_DIR, exist_ok=True)


# ============================================================
# 1. ПАРСИНГ ДОКУМЕНТОВ
# ============================================================

def extract_docx(filepath):
    """Извлекает текст из docx, сохраняя структуру параграфов"""
    doc = docx.Document(filepath)
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append({
                "text": text,
                "style": p.style.name if p.style else "Normal",
                "bold": any(run.bold for run in p.runs if run.bold),
            })

    # Extract tables too
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(rows)

    return paragraphs, tables


def parse_rules_document():
    """Парсинг 'Правила и традиции 2025'"""
    filepath = os.path.join(INPUT_DIR, "Правила_и_традиции_2025.docx")
    paragraphs, tables = extract_docx(filepath)

    # Структурируем правила по секциям
    rules = {
        "title": "Правила посещения детского центра «Про Дети»",
        "raw_paragraphs": [p["text"] for p in paragraphs],
        "sections": [],
    }

    current_section = {"title": "Общие правила", "items": []}
    for p in paragraphs:
        text = p["text"]
        # Detect section headers (bold or specific patterns)
        if p.get("bold") or text.startswith("ВНИМАНИЕ") or text.isupper():
            if current_section["items"]:
                rules["sections"].append(current_section)
            current_section = {"title": text, "items": []}
        else:
            current_section["items"].append(text)

    if current_section["items"]:
        rules["sections"].append(current_section)

    return rules


def parse_contract_document():
    """Парсинг Договора оферты"""
    filepath = os.path.join(INPUT_DIR, "Договор_оферты-_Близнин.docx")
    paragraphs, tables = extract_docx(filepath)

    contract = {
        "title": "Договор публичной оферты на оказание услуг",
        "entity": "ИП Близнин Семен Алексеевич",
        "inn": "290222039347",
        "city": "Северодвинск",
        "date": "01 сентября 2025",
        "raw_paragraphs": [p["text"] for p in paragraphs],
        "sections": [],
    }

    current_section = {"title": "Преамбула", "items": []}
    for p in paragraphs:
        text = p["text"]
        # Section detection: numbered sections like "1.", "2." etc
        section_match = re.match(r'^(\d+)\.\s+(.+)', text)
        if section_match and len(text) < 100:
            if current_section["items"]:
                contract["sections"].append(current_section)
            current_section = {"title": text, "items": []}
        else:
            current_section["items"].append(text)

    if current_section["items"]:
        contract["sections"].append(current_section)

    return contract


# ============================================================
# 2. ПАТТЕРНЫ ОТВЕТОВ ОПЕРАТОРОВ
# ============================================================

def analyze_operator_patterns(sessions):
    """Анализ паттернов ответов операторов из реальных диалогов"""

    # Интент-маппинг по ключевым словам
    intent_patterns = {
        "TRIAL_BOOKING": {
            "keywords": ["пробное", "пробный", "пробного", "первое занятие", "попробовать"],
            "examples": [],
        },
        "PRICING": {
            "keywords": ["цена", "стоимость", "абонемент", "сколько стоит", "прайс", "оплат"],
            "examples": [],
        },
        "SCHEDULE": {
            "keywords": ["расписание", "во сколько", "в какое время", "когда занятия", "график"],
            "examples": [],
        },
        "BOOKING": {
            "keywords": ["записать", "запись", "записаться", "свободн", "мест"],
            "examples": [],
        },
        "CANCELLATION": {
            "keywords": ["отменить", "не придем", "не сможем", "не будет", "пропуск", "пропустим"],
            "examples": [],
        },
        "MAKEUP_CLASS": {
            "keywords": ["отработк", "перенос", "перенести", "отработать"],
            "examples": [],
        },
        "AGE_GROUP": {
            "keywords": ["возраст", "лет", "года ребен", "месяц", "группа мини"],
            "examples": [],
        },
        "LOCATION": {
            "keywords": ["адрес", "где находит", "как добрат", "проехать", "центр"],
            "examples": [],
        },
        "FEEDBACK": {
            "keywords": ["спасибо", "понрави", "благодар", "отлично", "замечательно"],
            "examples": [],
        },
        "GENERAL_INFO": {
            "keywords": ["подскажите", "расскажите", "информац", "вопрос"],
            "examples": [],
        },
    }

    # Собираем пары вопрос-ответ
    qa_pairs = []
    for session in sessions:
        msgs = session["messages"]
        for i in range(len(msgs) - 1):
            if msgs[i]["role"] == "client" and msgs[i + 1]["role"] == "staff":
                client_text = msgs[i]["text"].strip()
                staff_text = msgs[i + 1]["text"].strip()

                if len(client_text) < 5 or len(staff_text) < 10:
                    continue

                # Classify intent
                intent = "OTHER"
                text_lower = client_text.lower()
                for intent_name, config in intent_patterns.items():
                    if any(kw in text_lower for kw in config["keywords"]):
                        intent = intent_name
                        break

                qa_pair = {
                    "session_id": session["session_id"],
                    "client_text": client_text,
                    "staff_text": staff_text,
                    "intent": intent,
                    "client_len": len(client_text),
                    "staff_len": len(staff_text),
                    "operator_id": msgs[i + 1].get("admin_author_id"),
                }
                qa_pairs.append(qa_pair)

                # Add to intent examples (limit 30 per intent)
                if intent != "OTHER" and len(intent_patterns[intent]["examples"]) < 30:
                    intent_patterns[intent]["examples"].append(qa_pair)

    # Analyze response templates/patterns
    greeting_patterns = Counter()
    closing_patterns = Counter()

    for pair in qa_pairs:
        staff = pair["staff_text"]
        # Greeting analysis
        first_line = staff.split("\n")[0].strip()
        if any(g in first_line.lower() for g in ["здравствуй", "добр", "привет"]):
            # Normalize greeting
            greeting = re.sub(r'[А-Я][а-я]+,\s*', '[ИМЯ], ', first_line)
            greeting_patterns[greeting] += 1

    return {
        "total_qa_pairs": len(qa_pairs),
        "intent_distribution": {
            intent: len(config["examples"])
            for intent, config in intent_patterns.items()
        },
        "intent_examples": {
            intent: config["examples"][:10]  # Top 10 per intent
            for intent, config in intent_patterns.items()
            if config["examples"]
        },
        "greeting_patterns": dict(greeting_patterns.most_common(10)),
        "all_qa_pairs": qa_pairs,
    }


# ============================================================
# 3. STYLE ANALYSIS — как операторы общаются
# ============================================================

def analyze_communication_style(qa_pairs):
    """Анализ стиля общения операторов"""

    emoji_usage = Counter()
    formality_indicators = {
        "formal_you": 0,  # Вы, Вас, Вам
        "informal_you": 0,  # ты, тебе
        "name_usage": 0,
        "emoji_count": 0,
    }

    emoji_pattern = re.compile(r'[\U0001F600-\U0001F64F\U0001F300-\U0001F5FF\U0001F680-\U0001F6FF\U0001F900-\U0001F9FF\U00002702-\U000027B0\U0001FA00-\U0001FA6F\U0001FA70-\U0001FAFF\U00002600-\U000026FF\U0000FE00-\U0000FE0F]')

    staff_texts = [p["staff_text"] for p in qa_pairs]

    for text in staff_texts:
        # Emoji
        emojis = emoji_pattern.findall(text)
        for e in emojis:
            emoji_usage[e] += 1
        formality_indicators["emoji_count"] += len(emojis)

        # Formal "Вы"
        if re.search(r'\b[Вв]ы\b|\b[Вв]ас\b|\b[Вв]ам\b', text):
            formality_indicators["formal_you"] += 1
        if re.search(r'\bты\b|\bтебе\b|\bтебя\b', text):
            formality_indicators["informal_you"] += 1

    # Average message patterns
    avg_staff_len = sum(len(t) for t in staff_texts) / len(staff_texts) if staff_texts else 0
    avg_sentences = sum(len(re.split(r'[.!?]+', t)) for t in staff_texts) / len(staff_texts) if staff_texts else 0

    style = {
        "formality": "formal" if formality_indicators["formal_you"] > formality_indicators["informal_you"] * 5 else "mixed",
        "formality_stats": formality_indicators,
        "avg_response_length_chars": round(avg_staff_len),
        "avg_sentences_per_response": round(avg_sentences, 1),
        "top_emojis": dict(emoji_usage.most_common(15)),
        "emoji_per_message": round(formality_indicators["emoji_count"] / len(staff_texts), 2) if staff_texts else 0,
        "style_notes": [],
    }

    # Generate style notes
    if style["formality"] == "formal":
        style["style_notes"].append("Операторы используют формальное обращение на 'Вы'")
    if style["emoji_per_message"] > 0.5:
        style["style_notes"].append("Активное использование эмодзи (в среднем {:.1f} на сообщение)".format(style["emoji_per_message"]))
    if avg_staff_len > 100:
        style["style_notes"].append("Развёрнутые ответы (в среднем {} символов)".format(round(avg_staff_len)))

    return style


# ============================================================
# MAIN
# ============================================================

def main():
    print("\n" + "=" * 70)
    print("ЭТАП 2: ИЗВЛЕЧЕНИЕ БАЗЫ ЗНАНИЙ И ПАТТЕРНОВ")
    print("=" * 70)

    # 1. Парсинг документов
    print("\n[1/4] Парсинг документов...")
    rules = parse_rules_document()
    contract = parse_contract_document()
    print(f"  Правила: {len(rules['raw_paragraphs'])} параграфов")
    print(f"  Договор: {len(contract['raw_paragraphs'])} параграфов, {len(contract['sections'])} секций")

    # 2. Загрузка сессий
    print("\n[2/4] Загрузка реконструированных сессий...")
    with open(os.path.join(OUTPUT_DIR, "sessions_full.json"), "r", encoding="utf-8") as f:
        sessions = json.load(f)
    print(f"  Загружено сессий: {len(sessions)}")

    # 3. Анализ паттернов
    print("\n[3/4] Анализ паттернов ответов операторов...")
    patterns = analyze_operator_patterns(sessions)
    print(f"  Всего QA-пар: {patterns['total_qa_pairs']}")
    print(f"  Распределение по интентам:")
    for intent, count in sorted(patterns["intent_distribution"].items(), key=lambda x: -x[1]):
        if count > 0:
            print(f"    {intent:>20}: {count}")

    # 4. Анализ стиля
    print("\n[4/4] Анализ стиля общения...")
    style = analyze_communication_style(patterns["all_qa_pairs"])
    print(f"  Формальность: {style['formality']}")
    print(f"  Средняя длина ответа: {style['avg_response_length_chars']} символов")
    print(f"  Эмодзи на сообщение: {style['emoji_per_message']}")
    print(f"  Топ эмодзи: {list(style['top_emojis'].keys())[:5]}")
    for note in style["style_notes"]:
        print(f"  → {note}")

    # === ЭКСПОРТ ===
    print("\n" + "-" * 70)
    print("ЭКСПОРТ БАЗЫ ЗНАНИЙ")
    print("-" * 70)

    # Knowledge Base
    knowledge_base = {
        "studio": {
            "name": "Детский центр «Про Дети»",
            "city": "Северодвинск",
            "owner": "ИП Близнин Семен Алексеевич",
            "inn": "290222039347",
        },
        "rules": rules,
        "contract": contract,
        "communication_style": style,
    }

    with open(os.path.join(KB_DIR, "knowledge_base.json"), "w", encoding="utf-8") as f:
        json.dump(knowledge_base, f, ensure_ascii=False, indent=2)
    print(f"  knowledge_base.json: основная база знаний")

    # Intent examples (for few-shot)
    with open(os.path.join(KB_DIR, "intent_examples.json"), "w", encoding="utf-8") as f:
        json.dump(patterns["intent_examples"], f, ensure_ascii=False, indent=2)
    print(f"  intent_examples.json: примеры по интентам ({sum(len(v) for v in patterns['intent_examples'].values())} примеров)")

    # All QA pairs (for RAG indexing later)
    with open(os.path.join(OUTPUT_DIR, "qa_pairs.json"), "w", encoding="utf-8") as f:
        json.dump(patterns["all_qa_pairs"], f, ensure_ascii=False, indent=2)
    print(f"  qa_pairs.json: {len(patterns['all_qa_pairs'])} QA-пар для индексации")

    # Communication style guide (human-readable)
    style_guide = generate_style_guide(style, patterns)
    with open(os.path.join(KB_DIR, "style_guide.md"), "w", encoding="utf-8") as f:
        f.write(style_guide)
    print(f"  style_guide.md: гайд по стилю общения")

    # Rules plain text (for prompts)
    with open(os.path.join(KB_DIR, "rules_plain.txt"), "w", encoding="utf-8") as f:
        f.write("\n\n".join(rules["raw_paragraphs"]))
    print(f"  rules_plain.txt: правила в текстовом формате")

    print(f"\n✅ Этап 2 завершён. База знаний в {KB_DIR}/")


def generate_style_guide(style, patterns):
    """Генерация текстового гайда по стилю"""
    guide = """# Гайд по стилю общения — Детский центр «Про Дети»
## Автоматически извлечён из {total} реальных QA-пар

## Общий стиль
- Обращение: на **{formality}**
- Средняя длина ответа: **{avg_len} символов** (~{sentences} предложений)
- Эмодзи: **{emoji_rate}** на сообщение

## Типичные приветствия
{greetings}

## Используемые эмодзи
{emojis}

## Заметки по стилю
{notes}

## Примеры лучших ответов по интентам

""".format(
        total=patterns["total_qa_pairs"],
        formality="Вы" if style["formality"] == "formal" else "смешанное",
        avg_len=style["avg_response_length_chars"],
        sentences=style["avg_sentences_per_response"],
        emoji_rate=style["emoji_per_message"],
        greetings="\n".join(f"- `{g}` ({c} раз)" for g, c in list(patterns["greeting_patterns"].items())[:5]),
        emojis=" ".join(f"{e}({c})" for e, c in list(style["top_emojis"].items())[:10]),
        notes="\n".join(f"- {n}" for n in style["style_notes"]),
    )

    # Add examples per intent
    for intent, examples in patterns["intent_examples"].items():
        if not examples:
            continue
        guide += f"\n### {intent}\n"
        for ex in examples[:3]:
            guide += f"\n**Клиент:** {ex['client_text'][:200]}\n"
            guide += f"**Оператор:** {ex['staff_text'][:300]}\n"

    return guide


if __name__ == "__main__":
    main()
